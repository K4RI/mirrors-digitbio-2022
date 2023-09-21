# -*- coding: utf-8 -*-
# Created on Mon Jun 23 2022
# @author: K4RI

""" Inférence de règles de causalité sur les données agronomiques du projet
MIRRORS, à partir de plusieurs algorithmes existants. """

from codes_sources.FormalCausalityAnalysis import FCaA
from codes_sources.origo_master3.store import Store
from codes_sources.origo_master3.score import score_ergo_origo_dc
from codes_sources.origo_master3.utils import concatenate
from codes_sources.origo_master3.multivariate_discretizer \
        import IPMvDiscretizer as MDiscretizer
# NMLMvDiscretizer, EquiWidthMvDiscretizer, \
# EquiFrequencyMvDiscretizer, IPMvDiscretizer
from codes_sources.origo_master3.univariate_discretizer \
        import UnivariateNMLDiscretizer as UDiscretizer
# UnivariateNMLDiscretizer, UnivariateEquiWidthDiscretizer, \
# UnivariateEquiFrequencyDiscretizer, UnivariateIPDiscretizer

import ast
from docx import Document
from docx.shared import RGBColor, Pt
import itertools
import matplotlib.pyplot as plt
import networkx as nx
import numpy as np
import os
import random
import shutil
import subprocess
import time


start_wd = os.getcwd()


def matcol(M: list, v: int):
    """ Renvoie la v-ième colonne de M."""
    return [row[v] for row in M]


def preprocess(version: int, verb: bool = False):
    """ Récupère le tableau Excel dans une matrice.

    Parameters
    ----------
    version : int
        1 ou 2, selon le tableur que l'on souhaite utiliser.
    verb: bool
        Cette fonction est appelée en tant qu'annexe dans d'autres parties
        du programme, il n'est pas nécessaire d'écrire dans la console
        à chaque appel.

    Returns
    ----------
    list
        La matrice des valeurs.
    list
        L'en-tête contenant les noms des attributs.
    """
    if verb:
        print("\nPreprocessing du tableur...")

    if version == 1:
        filename = r"RAPSODYN_Zachary_20220414\DATA\rapsodyn_temp_25_30.csv"
    else:
        filename = r"RAPSODYN_Zachary_20220414\DATA" \
                    "\Rapsodyn_data_20220519_v2.csv"

    if verb:
        print("Lecture de '%s'" % filename)

    with open(filename, "r") as file:  # on ouvre le tableur
        data = file.readlines()
        data = [x.replace('Ã©', 'e') for x in data]
        header, data = data[0].split("\n")[0].split(";")[2:], data[1:]
        file.close()

    M = []
    if version == 1:
        sep = ''
    else:
        sep = 'NA'
    for line in data:
        row = line.split("\n")[0].split(";")
        if row[-4] != sep:
            # on ne compte que les lignes complètes ici 118/170
            M.append(row[2:])

    # on remplace les valeurs catégorielles par des nombres
    climates = set(matcol(M, 0))
    dicoclimates = {v: i for (i, v) in enumerate(climates)}

    genotypes = set(matcol(M, 1))
    dicogenotypes = {v: i for (i, v) in enumerate(genotypes)}

    n, p = len(M[0]), len(M)
    for i in range(p):
        M[i][0] = dicoclimates[M[i][0]]
        M[i][1] = dicogenotypes[M[i][1]]
        # on évalue en float le reste des valeurs
        for j in range(2, n):
            if version == 1:
                M[i][j] = float(M[i][j].replace(',', '.'))
            else:
                M[i][j] = float(M[i][j])

    n, p = len(M[0]), len(M)
    # attention ! les valeurs toujours nulles posent souci pour les régressions
    # donc on les retire
    # ici P300_30 P600_30 pour la v1
    # et P300>_30 P600>_30 P1000_<5 harvest_<5 pour la v2
    for var in range(n-1, -1, -1):
        col = list(map(float, matcol(M, var)))
        if all(x == 0 for x in col):
            for row in M:
                del row[var]
            n -= 1
            del header[var]

    # print(M[0], header)
    if verb:
        print("...terminé !")
    return M, header


# ########## PARTIE I - CFCA ##########

def CFCA_process(version: int):
    """ Exécuter l'algorithme d'Alexandre sur nos données.
        Attention : dure plusieurs heures. """
    M, header = preprocess(version)
    print("\nPart. I - Lancement de l'algorithme CFCA...")

    FCAdir = "codes_sources/FormalCausalityAnalysis/"

    if version == 1:
        mat_file, file = "matMFCaA.txt", "file25_30.txt"
    else:
        mat_file, file = "matMFCaA2.txt", "file25_302.txt"
    print("Lecture de '%s'" % mat_file)

    # on prépare les fichiers
    with open(FCAdir + mat_file, "w") as f:
        f.write(str(M))
        f.close()

    FCaA.matrix2File(M, FCAdir + file)

    # puis on lance le programme
    subprocess.call(["py", "FCaA.py", mat_file, file],
                    cwd='codes_sources/FormalCausalityAnalysis')

    outputname = 'Expe_agro_Zac_v%i.txt' % version
    for f in [mat_file, file, outputname]:
        shutil.move(FCAdir + f, 'files/1_cfca/' + f)

    print("Règles inférées écrites dans 'files/1_cfca/%s'" % outputname)
    print("...terminé !")


def CFCA_postprocess(version: int, dimY: int):
    """ Trie les règles inférées, en ne gardant que
        les causes exclusivement climat et les effets plante.

    Parameters
    ----------
    dimY : int
        Le nombre de colonnes des variables plante.
    """
    M, header = preprocess(version)
    print("\nPart. I - Tri des règles d'inférence...")

    outputname = "files/1_cfca/Expe_agro_Zac_v%i.txt" % version
    print("Lecture de '%s'" % outputname)
    # on lit le fichier de sortie généré dans CFCA_process()
    with open(outputname, "r") as file:
        data = file.readlines()
        file.close()

    print('%i règles initiales' % (len(data)-1))

    seuil = len(header)-dimY
    outputname2 = outputname[:-4] + ' TRIÉ.txt'
    cpt = 0
    with open(outputname2, "w") as fout:
        for line in data[1:]:
            # on trouve les causes et effets de chaque règle
            causes, effets = line.split(' -> ')
            causes = list(map(int, causes.split(' ')))
            effets = effets.split(' \n')[0]
            effets = list(map(int, effets.split(' \n')[0].split(' ')))
            effets = [e for e in effets if e >= seuil]
            if len(effets) > 0 and not any([c >= seuil for c in causes]):
                # si la règle nous convient, on l'écrit
                cpt += 1
                for c in causes:
                    fout.write(str(c))
                    fout.write(' ')
                fout.write('-> ')
                for e in effets:
                    fout.write(str(e))
                    fout.write(' ')
                fout.write('\n')
        fout.close()

    print('%i règles retenues' % cpt)
    print("Écrites dans '%s'" % outputname2)

    print("...terminé !")


# ########## PARTIE IIa - BicPAMS ##########

def bicpams_preprocess(version: int):
    """ Remodèle le tableau vers une forme acceptée par BicPAMS. """
    M, header = preprocess(version)
    print("\nPart. IIa - Pre-processing pour BicPAMS...")
    if version == 1:
        bicfile = "files/2a_bicpams/bicpams25_30.txt"
    else:
        bicfile = "files/2a_bicpams/bicpams25_302.txt"

    f = open(bicfile, "w")
    f.write(',')
    f.write(','.join(header))
    i = 1
    for row in M:
        f.write('\n')
        f.write(str(i))
        f.write(',')
        f.write(','.join(list(map(str, row))))
        i += 1
    f.close()
    print("Sauvegardé dans '%s'" % bicfile)

    print("...terminé !")


# hélas pas de version exécutable en console, seulement un IDE
# Normalization : Column, Noise handler: Multitem
# donne des résultats en constant et en additive


def biclusterOK(b: list, climat: list, plants: list):
    """ Fonction annexe, trie les biclusters.

    Parameters
    ----------
    b : list
        Liste des biclusters.
    climat : list
        Liste des variables climat.
    plants : list
        Liste des variables plante.

    Returns
    -------
    list
        Les biclusters contenant à la fois des
        variables climat et des variables plante.
    """
    names = b['Y']
    return any([x in plants for x in names]) \
        and any([x in climat for x in names])


def biclustersdupl(bs: list):
    """ Retire les doublons d'un ensemble de biclusters.
    Parameters
    ----------
    bs : list
        Liste des biclusters.

    Returns
    -------
    list
        La même liste sans doublons.
    """
    b2 = []
    ly = []
    for b in bs:
        if b['Y'] not in ly:
            b2.append(b)
            ly.append(b['Y'])
    return b2


def bicpams_postprocess(version: int, coh: str, dimY: int):
    """ Récupère les biclusters du fichier sorti par BicPAMS,
        puis trie similairement à CFCA_postprocess().

    Parameters
    ----------
    coh : str
        add ou cst, selon le fichier qu'on souhaite analyser.
    """
    M, header = preprocess(version)
    print("\nPart. IIa - Récupération et tri des biclusters...")
    file_clusters = 'files/2a_bicpams/result_bicpams_%s_v%i.htm' \
                    % (coh, version)

    plants, climat = header[:-dimY], header[-dimY:]
    biclusters = []
    print("Lecture de '%s'" % file_clusters)
    with open(file_clusters, "r") as f:
        # on parse le fichier .htm
        data = f.read().split('<tr><td><b>')
        for line in data[1:]:
            size, line = line.split('</b></td><td><b>X</b>=')
            X, line = line.split('<br><b>Y</b>=')
            Y = line.split('<br><br></td></tr>')[0]
            Y = Y.replace("[", "['").replace(", ", "', '").replace("]", "']")
            # on consigne les biclusters dans une liste
            biclusters.append({'size': ast.literal_eval(size),
                               'X': ast.literal_eval(X),
                               'Y': ast.literal_eval(Y)})
        f.close()

    print('%i biclusters inférés' % len(biclusters))
    biclusters = biclustersdupl(biclusters)
    print('%i biclusters non-doublons' % len(biclusters))
    # on ne garde que les biclusters cohérents
    biclustersOK = list(filter(
                        lambda b: biclusterOK(b, climat, plants), biclusters))
    print('%i biclusters retenus' % len(biclustersOK))

    # on les écrit joliment dans un fichier .docx
    file_clusters_out = file_clusters[:-4] + '_TRIÉ.docx'
    red, black, green = RGBColor(255, 0, 0), \
        RGBColor(0, 0, 0), \
        RGBColor(0, 255, 0)
    document = Document()
    p = document.add_paragraph()
    for b in biclustersOK:
        run = p.add_run('[')
        run.font.color.rgb = black
        for x in b['Y'][:-1]:
            run = p.add_run("'%s', " % x)
            if x in climat:
                run.font.color.rgb = green
            else:
                run.font.color.rgb = red
        run = p.add_run("'%s'" % b['Y'][-1])
        run.font.color.rgb = green
        run = p.add_run('] ')
        run.font.color.rgb = black
        p.add_run(str(b['size'])).bold = True
        run.font.color.rgb = black
        run = p.add_run().add_break()

    document.save(file_clusters_out)

    print("Sauvegardé dans '%s'" % file_clusters_out)

    # for b in biclustersOK: print(b['Y'], b['size'])

    print("...terminé !")


# ########## PARTIE IIb - Pypingen ##########

def limsup(n: int, lims: list):
    for x in lims:
        if n < x:
            return x
    return n


def pypingen_binarisation(version: int):
    """ Binarise le tableau v2.
        Les nouveaux indices sont "i->j", et x[i->j] = sgn(x[i]-x[j]). """
    if version == 1:
        return  # pour la v1, le fichier binarisé existe déjà

    print("\nPart. IIb - Binarisation du tableur...")

    filename = r"RAPSODYN_Zachary_20220414\DATA\Rapsodyn_data_20220519_v2.csv"
    file_bin = r"RAPSODYN_Zachary_20220414\DATA\rapsodyn_binary_v2.csv"

    print("Lecture de '%s'" % filename)
    with open(filename, "r") as file:
        data = file.readlines()
        data = [x.replace('Ã©', 'e') for x in data]
        file.close()
    header, data = data[0].split("\n")[0].split(";")[4:], data[1:]

    n = len(data)
    p = len(header)
    for i in range(n):
        data[i] = data[i].split("\n")[0].split(";")[4:]

    limites = [6, 12, 18, 30, 42, 66, 68, 70, 74, 78, 81, 91, 93, 95, 97, 99,
               103, 107, 111, 129, 134, 136, 138, 140, 144, 146, 148, 156,
               160, 164, 166, 168, 169]

    # on crée un nouveau classeur
    with open(file_bin, "w") as f:
        f.write(';' + ';'.join(header) + '\n')
        for i in range(n):
            for j in range(limsup(i, limites), n):
                f.write('%i->%i' % (i, j))
                for k in range(p):
                    f.write(';')
                    if data[i][k] == 'NA' or data[j][k] == 'NA':
                        pass
                    # selon la valeur relative de x[i] et x[j],
                    # on note -1 0 ou 1
                    elif data[i][k] < data[j][k]:
                        f.write('-1,0')
                    elif data[i][k] > data[j][k]:
                        f.write('1,0')
                    elif data[i][k] == data[j][k]:
                        f.write('0,0')
                f.write('\n')
        f.close()

    print("Sauvegardé dans '%s'" % file_bin)

    print("...terminé !")


def pypingen_plusoumoins(version: int):
    """ Traite le tableau binaire précédent :
        changer les 1 en "+", et les -1 en "-". """

    print("\nPart. IIb - Traduction du tableur binarisé en + - ...")

    if version == 1:
        file_bin = r"RAPSODYN_Zachary_20220414\DATA\rapsodyn_binary.csv"
        file_plusmoins = "files/2b_pypingen/pypingen_binary.txt"
    else:
        file_bin = r"RAPSODYN_Zachary_20220414\DATA\rapsodyn_binary_v2.csv"
        file_plusmoins = "files/2b_pypingen/pypingen_binary_v2.txt"

    print("Lecture de '%s'" % file_bin)
    with open(file_bin, "r") as fin:
        data = fin.readlines()[1:]
        fin.close()
    n, p = len(data), data[0].count(';')

    with open(file_plusmoins, "w") as fout:
        for i in range(n):
            row = data[i].split("\n")[0].split(";")[1:]
            for j in range(p):
                if row[j] == '-1,0':
                    fout.write('-')
                elif row[j] == '1,0':
                    fout.write('+')
                else:
                    fout.write('+')  # random.choice(['+', '-'])) # le hasard ?
                if j < p-1:
                    fout.write(' ')
            if i < n-1:
                fout.write('\n')
        fout.close()

    print("Sauvegardé dans '%s'" % file_plusmoins)

    print("...terminé !")

# amélioration : inclure un signe '?' dans Pypingen
# pour empêcher toute inclusion à un cluster
# que faire des cases vides ?


def pypingen_process(version: int):
    """ Lance l'algorithme Pypingen. Attention : très très long. """

    print("\nPart. IIb - Lancement de l'algorithme Pypingen'...")

    if version == 1:
        file = "pypingen_binary.txt"
    else:
        file = "pypingen_binary_v2.txt"

    pypingenDIR = 'codes_sources/CSC/'
    shutil.move('files/2b_pypingen/' + file, pypingenDIR + file)

    print("Lecture de '%s'" % file)
    subprocess.call(["py", "addIntent.py", pypingenDIR + file],
                    cwd=pypingenDIR)

    outputfile = "addintent_%s_c1_r1_lattice.txt" % file
    shutil.move(pypingenDIR + outputfile, 'files/2b_pypingen/' + outputfile)
    shutil.move(pypingenDIR + file, 'files/2b_pypingen/' + file)

    print("Sauvegardé dans 'files/2b_pypingen/%s'" % outputfile)

    print("...terminé !")


# ########## PARTIE IIc - Coron ##########

def triplet_cat(x, dico):
    l = [0]*len(dico)
    l[dico[x]] = 1
    return l


def coron_preprocess(version: int, verb: bool = False):
    """ Binarise le tableau : chaque attribut numérique est divisé
    en deux binaires : 'x[i]<mediane(x)' et 'x[i]>=mediane(x)'.

    Returns
    ----------
    list
        La matrice des valeurs binaires.
    list
        L'en-tête contenant les noms des attributs binarisés.
    list
        Le nombre de valeurs possibles pour chaque attribut catégoriel.
    """

    if verb:
        print("\nPart IIc - Preprocessing binaire du tableur...")

    if version == 1:
        filename = r"RAPSODYN_Zachary_20220414\DATA\rapsodyn_temp_25_30.csv"
    else:
        filename = r"RAPSODYN_Zachary_20220414\DATA" \
                    "\Rapsodyn_data_20220519_v2.csv"

    if verb:
        print("Lecture de '%s'" % filename)

    with open(filename, "r") as file:
        data = file.readlines()
        data = [x.replace('Ã©', 'e') for x in data]
        header, data = data[0].split("\n")[0].split(";")[2:], data[1:]
        file.close()
    M = []

    if version == 1:
        sep = ''
    else:
        sep = 'NA'

    for line in data:
        if version == 1:
            row = line.replace(',', '.').split("\n")[0].split(";")
        else:
            row = line.split("\n")[0].split(";")
        M.append(row[2:])
    n, p = len(M[0]), len(M)

    # on retire les attributs toujours nuls
    for var in range(n-1, -1, -1):
        col = matcol(M, var)
        if all(x == '0' for x in col):
            for row in M:
                del row[var]
            n -= 1
            del header[var]

    lims = []
    header2 = []
    lcat = []
    for var in range(n):
        col = matcol(M, var)
        try:  # si la variable est numérique (type int ou float)
            L = [ast.literal_eval(x) for x in col if x != sep]
            # selon si la valeur est en-dessous ou au-dessus de la médiane...
            lims.append(np.median(L))
            # ... on associera un des deux attributs : LOW et HIGH
            header2.append('%s_LOW' % header[var])
            header2.append('%s_HIGH' % header[var])
        except:  # si la variable est catégorielle (type str)
            s = set(col)
            dico = {v: i for (i, v) in enumerate(s)}
            lims.append(dico)
            header2.extend(list(s))
            # on associera l'attribut de la catégorie
            lcat.append(len(dico))

    M2 = []
    for i in range(p):
        L = []
        for var in range(n):
            if isinstance(lims[var], dict):  # var catégorielle, str
                L.extend(triplet_cat(M[i][var], lims[var]))
            else:  # var catégorielle, int ou float
                if M[i][var] == sep:
                    L.extend([0, 0])
                elif float(M[i][var]) <= lims[var]:
                    # si basse valeur : LOW <- 1, HIGH <- 0
                    L.extend([1, 0])
                else:
                    # si haute valeur : LOW <- 0, HIGH <- 1
                    L.extend([0, 1])
        M2.append(L)

    if verb:
        print("...terminé !")

    return M2, header2, lcat


def coron_binarisation(version: int):
    """ Transcrit les valeurs binarisées dans un fichier. """

    print("\nPart IIc - Binarisation du tableur...")

    M2, header2, _ = coron_preprocess(version)
    with open('files/2c_coron/binary_new_v%i.csv' % version, 'w') as f:
        f.write(';'.join(header2))
        for row in M2:
            f.write('\n')
            f.write(';'.join(map(str, row)))
        f.close()

    print("Sauvegardé dans 'files/2c_coron/binary_new_v%i.csv'" % version)

    print("...terminé !")


def coron_mat2File(version: int):
    """ Transforme la matrice en fichier compatible avec Coron. """

    print("\nPart IIc - Transformation en fichier Coron...")
    print("Lecture de 'files/2c_coron/binary_new_v%i.csv'" % version)

    with open('files/2c_coron/binary_new_v%i.csv' % version, "r") as file:
        data = file.readlines()
        header, data = data[0].split("\n")[0].split(";"), data[1:]
        file.close()

    n = len(data)
    for i in range(n):
        data[i] = data[i].split('\n')[0].split(';')

    with open('files/2c_coron/rapsodynCoron_v%i.rcf'
              % version, "w") as f:
        f.write('# rapsodynCoron_v%i\n\n' % version)
        f.write("""[Relational Context]
Default Name
[Binary Relation]
Name_of_dataset
""")
        f.write(' | '.join(map(str, range(1, n+1))))
        f.write('\n')
        f.write(' | '.join(header))
        f.write('\n')
        for row in data:
            f.write(' '.join(row))
            f.write('\n')
        f.write('[END Relational Context]')
        f.close()
    print("Sauvegardé dans 'files/2c_coron/rapsodynCoron_v%i.rcf'" % version)

    print("...terminé !")


def coron_process(version: int, min_supp: int, min_conf: int):
    """ Lance Coron sur notre fichier.

    Parameters
    ----------
    min_supp : int
        Support. Pourcentage minimal d'objets devant avoir X et Y.
    min_conf : int
        Confidence. Pourcentage minimal d'objets de X devant avoir Y.
    """

    print("\nPart IIc - Exécution de Coron...")

    file = 'rapsodynCoron_v%i.rcf' % version
    print("Lecture de '%s'" % file)

    coronDIR = "codes_sources/coron-0.8/"
    shutil.move('files/2c_coron/' + file, coronDIR + file)

    outputfile = "coron%i.log" % version
    subprocess.call(["sh", "core02_assrulexMODIF.sh",
                     "rapsodynCoron_v%i.rcf" % version, str(min_supp) + '%',
                     str(min_conf) + '%', "-alg:close", "-rule:closed",
                     "-of:%s" % outputfile],
                    cwd=coronDIR)
    # sh core02_assrulexMODIF.sh rapsodynCoron_v1.rcf 50% 50% -alg:close
    # -rule:closed -of:coron1.log

    shutil.move(coronDIR + file, 'files/2c_coron/' + file)
    shutil.move(coronDIR + outputfile, 'files/2c_coron/' + outputfile)

    with open('files/2c_coron/coron%i.log' % version, 'r') as f:
        cpt = f.read().count('\n{')
        f.close()
    print("%i règles extraites" % cpt)

    print("Sauvegardé dans 'files/2c_coron/coron%i.log'" % version)

    print("...terminé !")


def backrules(x, lcat):
    """ Retranscrit des attributs binarisés vers les attributs initiaux.

    Parameters
    ----------
    x : int
        Un attribut binarisé.
    lcat : list
        Le nombre de valeurs possibles pour chaque attribut catégoriel.

    Return
    ----------
    int
        L'attribut initial correspondant.
    """

    lim = 0
    n = len(lcat)
    for i in range(n):
        lim += lcat[i]
        if x < lim:
            return i
    return n + ((x - lim) // 2)


def coron_postprocess(version, dimY):
    """ Extrait et trie les règles d'association. """

    print("\nPart IIc - Tri des règles...")

    _, header, lcat = coron_preprocess(version)

    inputfile = "coron%i.log" % version
    outputfileNames = "coron_ruleNames_v%i.txt" % version
    outputfile = "coron_rules_v%i.txt" % version

    print("Lecture de 'files/2c_coron/%s'" % inputfile)
    coronDIR = "codes_sources/coron-0.8/"
    shutil.move('files/2c_coron/' + inputfile, coronDIR + inputfile)

    # on garde les règles n'ayant que des attributs plante à gauche
    n = len(header)
    indsY = '"' + '|'.join([str(x) for x in range(n-2*dimY+1, n+1)]) + '"'
    print(indsY)
    subprocess.call(["sh", "post01_filterRulesMODIF.sh",
                     inputfile, indsY, "-delete", "-left",
                     "-of:%s" % outputfile],
                    cwd=coronDIR)
    # sh post01_filterRulesMODIF.sh coron1.log "24|25|26|27|28|29|30|31|32|33"
    # -delete -left -of:coron_rules_v1.txt

    shutil.move(coronDIR + inputfile, 'files/2c_coron/' + inputfile)
    shutil.move(coronDIR + outputfile, 'files/2c_coron/' + outputfile)

    with open('files/2c_coron/%s' % outputfile, 'r') as f:
        data = f.readlines()
        f.close()

    rules = []
    ruleNames = []
    cpt = 0
    for row in data:
        if row[0] == '{':
            cpt += 1
            r, kw = row.split(' (')
            X, Y = r.split('} => {')
            X = X[1:].split(', ')
            Y = Y[:-1].split(', ')
            # ATTENTION : Coron débute l'indexation à 1, il faut tout décaler
            X = [int(x)-1 for x in X]
            Y = [int(y)-1 for y in Y if int(y) > n-2*dimY]

            kw = kw[5:].split('%]; suppL')[0]
            supp, _, conf = kw.split(' [')

            if Y:
                X1 = [str(backrules(x, lcat)) for x in X]
                Y1 = [str(backrules(y, lcat)) for y in Y]
                rules.append((X1, Y1, supp, conf))
                X = [header[x] for x in X]
                Y = [header[y] for y in Y]
                ruleNames.append((X, Y))
    print("%i règles après tri gauche" % cpt)

    with open('files/2c_coron/%s' % outputfileNames, "w") as f:
        for (X, Y) in ruleNames:
            f.write('%s -> %s\n' % (' '.join(X), ' '.join(Y)))
        f.close()

    with open('files/2c_coron/%s' % outputfile, "w") as f:
        for (X, Y, s, c) in rules:
            f.write('%s -> %s'
                    ' / %s / %s\n' % (' '.join(X), ' '.join(Y), s, c))
        f.close()

    print("%i règles après tri droit" % len(rules))
    print("Sauvegardé dans"
          "'files/2c_coron/coron_rules_v%i.txt'" % version)
    print("Noms sauvegardés dans"
          "'files/2c_coron/coron_ruleNames_v%i.txt'" % version)

    print("...terminé !")

# from graphviz import Source
# s = Source.from_file('tmp/trie.dot')
# s.view()
# # fonctionne avec v1 mais pas v2 (erreur de syntaxe)


# ########## PARTIE III - LatViz ##########

# créer un context.json à partir des biclusters binaires de BicPAMS ?


# ########## PARTIE IV - Globe ##########

def globe_preprocess(version: int, catcols: int):
    """ Similaire à preprocess(),
    mais supprime les variables catégorielles au lieu de les numériser.

    Parameters
    ----------
    catcols : int
        Nombre de colonnes catégorielles.
    """

    print("\nPart. IV - Pre-processing du tableur pour Globe...")

    if version == 1:
        filename = r"RAPSODYN_Zachary_20220414\DATA\rapsodyn_temp_25_30.csv"
        globefile = "files/4_globe/globe25_30.txt"
    else:
        filename = r"RAPSODYN_Zachary_20220414\DATA\Rapsodyn" \
                    "_data_20220519_v2.csv"
        globefile = "files/4_globe/globe25_302.txt"

    print("Lecture de '%s'" % filename)
    with open(filename, "r") as fin:
        data = fin.readlines()
        data = [x.replace('Ã©', 'e') for x in data]
        header, data = data[0].split("\n")[0].split(";")[2+catcols:], data[1:]
        fin.close()

    M = []
    n = data[0].count(';')-catcols-1
    if version == 1:
        sep = ''
    else:
        sep = 'NA'
    for line in data:
        if version == 1:
            line = line.replace(',', '.')
        # on ne compte pas les premières colonnes catégorielles
        row = line.split("\n")[0].split(";")[2+catcols:]
        if row[-4] != sep:
            M.append(row)

    p = len(M)
    for var in range(n-1, -1, -1):
        col = list(map(float, matcol(M, var)))
        if all(x == 0 for x in col):
            for row in M:
                del row[var]
            n -= 1
            del header[var]

    with open(globefile, "w") as fout:
        fout.write(str(n))
        fout.write('\n')
        for row in M:
            fout.write(','.join(row))
            fout.write('\n')
        fout.close()

    print("Sauvegardé dans '%s'" % globefile)

    print("...terminé !")


def globe_process(version: int):
    """ Lance l'algorithme Globe.
    Les détails d'exécution sont dans /logs. """

    print("\nPart. IV - Lancement de l'algorithme Globe...")

    if version == 1:
        file = "globe25_30.txt"
    else:
        file = "globe25_302.txt"

    print("Lecture de 'files/4_globe/%s'" % file)
    globeDIR = "codes_sources/globe_v20220213/"
    shutil.move("files/4_globe/" + file, globeDIR + 'data/' + file)

    output = subprocess.check_output(["py", "main.py", 'data/' + file],
                                     cwd=globeDIR)

    with open('files/4_globe/globe_output_v%i.log' % version, 'w') as f:
        f.write(output.decode("utf-8"))
        f.close()
    shutil.move(globeDIR + 'data/' + file, 'files/4_globe/' + file)

    print("Sauvegardé dans 'files/4_globe/globe_output_v%i.log'" % version)

    print("...terminé !")


def globe_postprocess(version: int):
    """ Afficher le diagramme causal obtenu par Globe. """

    M, header = preprocess(version)
    print("\nPart. IV - Affichage du diagramme causal de Globe...")

    print("Lecture de 'files/4_globe/globe_output_v%i.log'" % version)
    with open('files/4_globe/globe_output_v%i.log' % version, 'r') as f:
        output = f.read()
        f.close()

    # on extrait la matrice d'adjacence dans le fichier sorti par Globe
    if version == 1:
        searchedfile = "data/globe25_30.txt"
    else:
        searchedfile = "data/globe25_302.txt"
    M = output.split('END LOGGING FOR FILE: %s' % searchedfile)[1] \
              .replace('\n\n\n', '').replace('.]', '],') \
              .replace('.', ',').replace('],]', ']]')
    adj_matrix = ast.literal_eval(M)

    # on l'affiche visuellement
    plt.figure(1)
    G = nx.from_numpy_matrix(np.array(adj_matrix), create_using=nx.DiGraph)

    labels = {}
    n = len(adj_matrix)
    for k in range(n):
        labels[k] = '%i- %s' % (k, header[k+2])
        print(labels[k])
    layout = nx.planar_layout(G)

    options = {"edgecolors": "tab:gray", "node_size": 400, "alpha": 0.9}
    # les noeuds de variables climat en rouge, et variables plante en vert
    iEnd = len(header)-dimY-catcols
    nx.draw_networkx_nodes(G, layout, nodelist=range(iEnd),
                           node_color="tab:red", **options)
    nx.draw_networkx_nodes(G, layout,  nodelist=range(iEnd, iEnd+dimY),
                           node_color="tab:green", **options)
    nx.draw_networkx_edges(G, layout, width=1.0, alpha=0.5)
    nx.draw_networkx_labels(G, layout, labels, font_size=6,
                            font_color="black")
    plt.savefig('files/4_globe/globefig_v%i.png' % version)
    print("Image sauvegardée dans 'files/4_globe/globefig_v%i.png'" % version)
    plt.close()

    print("...terminé !")


# ########## PARTIE V - Origo ##########

def scores(res: list):
    """ Fonction annexe, affiche les résultats des algorithmes. """

    wins = res.count(1)
    tot = res.count(-1)+res.count(1)
    if tot:
        r = wins/tot
        return (wins, tot, 100*r)
    return (0, 0, 0)


def ruleName(header: list, iX: list, iY: list):
    """ Représente la règle d'inférence 'X->Y' comme chaîne de caractères.

    Parameters
    ----------
    header : list
        En-tête des attributs.
    iX : list
        Indices cause.
    iY : list
        Indices effet.

    Returns
    ----------
    str
        Représentation de la règle sous la forme 'X1,...,Xn -> Y1,...,Ym'
    """
    return '%s -> %s' % (', '.join([header[x] for x in iX]),
                         ', '.join([header[y] for y in iY]))


def kwName(k):
    """ Représente les indicateurs comme chaîne de caractères. """
    if isinstance(k, int):
        return 'size=%i, ' % k
    elif isinstance(k, tuple):
        return 'supp=%i, conf=%.2f%%, ' % k
    else:
        return ''


def origo(version: list, source: tuple, dimY: int, eps: float = 1e-3):
    """ Origo est un algorithme qui infère un sens entre
        deux attributs multivariées (= ensemble de plusieurs attributs).
        Vérifie qu'Origo infère le même sens que les méthodes 1 (CFCA)
        ou 2a (BicPAMS) dans les règles qu'on y a extraites.

    Parameters
    ----------
    source : tuple ou int
        source[0] est l'approche d'où proviennent les règles analysées.
        source[1] est un argument supplémentaire :
            pour source[0]=='CFCA':
                si les règles ont été triées par CFCA_postprocess()
            pour source[1]=='bicpams':
                'add' ou 'cst' selon le type de cohérence.
        / ou /
        source=='coron'
    eps : float
        Le seuil de validation par Origo.
    """

    M, header = preprocess(version)
    M = np.array(M)
    print("\nPart. V - Vérification par Origo des règles inférées...")
    os.chdir(start_wd)

    indsX, indsY, kwarg = [], [], []
    input = ''

    if (source[0] == 'CFCA' or source == 'coron'):
        # si on s'intéresse aux résultats de la méthode 1 ou 2c
        if source == 'coron':
            input = "files/2c_coron/coron_rules_v%i.txt" % version
        if source[0] == 'CFCA':
            input = "files/1_cfca/Expe_agro_Zac_" \
                    "v%i%s.txt" % (version, source[1]*' TRIÉ')
        print("Lecture de '%s'" % input)

        with open(input, 'r') as f:
            data = f.readlines()
            # on enlève l'en-tête si besoin
            try:
                if data[0][-8:] == 'seconds\n':
                    data = data[1:]
            except:
                pass
            f.close()

        for inf in data:
            # on extrait les indices de chaque règle
            rule = inf.split(' \n')[0]
            ix, iy = rule.split(' -> ')

            if source == 'coron':
                iy, s, c = iy.split(' / ')
                kwarg.append((int(s), float(c)))
            else:
                kwarg.append('')

            indsX.append(list(map(int, ix.split(' '))))
            indsY.append(list(map(int, iy.split(' '))))

    elif source[0] == 'bicpams':
        # si on s'intéresse aux résultats de la méthode 2a
        ext1 = {'add': 'additive', 'cst': 'constant'}
        input = 'files/2a_bicpams/result_bicpams_%s_v%i_' \
                'TRIÉ.docx' % (ext1[source[1]], version)
        print("Lecture de '%s'" % input)
        doc = Document(input)

        # le fichier de BicPAMS contient les noms de variables,
        # or pour la suite on a besoin des indices
        inds = {header[k]: k for k in range(len(header))}
        plants, climat = header[:-dimY], header[-dimY:]

        paras = doc.paragraphs[0].text.split('\n')[:-1]
        for para in paras:
            # à chaque ligne du fichier .docx,
            # on extrait variables climat et plante
            bicl, taille = para.split(" (")
            bicl = ast.literal_eval(bicl)
            taille = ast.literal_eval(taille.split(", ")[1][:-1])
            iX = [inds[x] for x in bicl if x in plants]
            iY = [inds[y] for y in bicl if y in climat]
            kwarg.append(taille)
            indsX.append(iX)
            indsY.append(iY)

    else:
        raise Exception("Merci d'entrer 'CFCA', 'bicpams' "
                        "ou 'coron' en source.")

    os.chdir('codes_sources\origo_master3')

    res_sign = {0: '?', -1: '✗', 1: '✓'}

    n = len(indsX)
    print('Nombre de règles :', n)
    res_ergo, res_origo, res_dc = [], [], []
    diffs_ergo, diffs_origo, diffs_dc = [], [], []
    cpt = 0
    for iX, iY, k in zip(indsX, indsY, kwarg):
        cpt += 1
        rule = '%s -> %s' % (' '.join(map(str, iX)), ' '.join(map(str, iY)))
        # on affiche le nom et indices de chaque variable de la règle
        print('\n----------%i/%i----------' % (cpt, n))
        print(rule)
        print(ruleName(header, iX, iY))
        if not isinstance(k, str):  # bicPAMS ou coron
            print('(%s)' % kwName(k)[:-2])
        print()

        # pour chacune des variables, on a ses valeurs dans le classeur
        X = [M[:, j] for j in iX]
        Y = [M[:, j] for j in iY]
        # mise en forme
        discretizer = MDiscretizer(X, Y)
        aX, Xd, aY, Yd = discretizer.discretize()
        aX = list(itertools.chain(*aX))
        aY = list(itertools.chain(*aY))
        rows = concatenate(Xd, Yd)
        store = Store(None, rows, aX + aY)

        ergo, origo, dc = score_ergo_origo_dc(aX, aY, store)
        # les "complexités de Kolmogorov" des deux sens pour chaque algo
        # ΔX→Y et ΔY→X
        diffs_ergo.append(ergo.Y_to_X - ergo.X_to_Y)
        diffs_origo.append(origo.Y_to_X - origo.X_to_Y)
        diffs_dc.append(dc.Y_to_X - dc.X_to_Y)

        # si ΔY→X > ΔX→Y, alors X→Y est inférée : c'est bon
        if diffs_ergo[-1] > eps:
            res_ergo.append(1)
        # si ΔY→X < ΔX→Y, alors Y→X est inférée : c'est pas bon
        elif diffs_ergo[-1] < -eps:
            res_ergo.append(-1)
        # si ΔY→X = ΔX→Y, on ne peut pas décider : le résultat est ignoré
        else:
            res_ergo.append(0)

        if diffs_origo[-1] > 0:
            res_origo.append(1)
        elif diffs_origo[-1] < 0:
            res_origo.append(-1)
        else:
            res_origo.append(0)

        if diffs_dc[-1] > 0:
            res_dc.append(1)
        elif diffs_dc[-1] < 0:
            res_dc.append(-1)
        else:
            res_dc.append(0)

        # affichage de la sortie de chaque algorithme
        print("ERGO : %s %.4f" % (res_sign[res_ergo[-1]], diffs_ergo[-1]))
        print("ORIGO : %s %.4f" % (res_sign[res_origo[-1]], diffs_origo[-1]))
        print("DC : %s %.4f" % (res_sign[res_dc[-1]], diffs_dc[-1]))

    # affichage de la proportion de règles correctement inférées
    bilan = '\n\n\n--------------------\n--------------------'
    bilan += '\nSCORE ERGO : %i/%i (%.1f%%)' % scores(res_ergo)
    bilan += '\nSCORE ORIGO : %i/%i (%.1f%%)' % scores(res_origo)
    bilan += '\nSCORE DC : %i/%i (%.1f%%)' % scores(res_dc)
    print(bilan)

    os.chdir(start_wd)

    file_out = 'files/5_origo/' + input.split('/')[-1].split('.')[0] \
               + '_origo.txt'
    ies = [i for i in range(n) if res_ergo[i] > 0]
    ios = [i for i in range(n) if res_origo[i] > 0]
    ids = [i for i in range(n) if res_dc[i] > 0]
    iS = [i for i in range(n) if
          (res_ergo[i] > 0 and res_origo[i] > 0 and res_dc[i] > 0)]

    # if isinstance(kwarg[0], int):  # bicPAMS, on trie par support
    #     ies.sort(key=lambda i: kwarg[i], reverse=True)
    #     ios.sort(key=lambda i: kwarg[i], reverse=True)
    #     ids.sort(key=lambda i: kwarg[i], reverse=True)
    #     iS.sort(key=lambda i: kwarg[i], reverse=True)
    # elif isinstance(kwarg[0], tuple):  # coron, on trie par confidence
    #     ies.sort(key=lambda i: kwarg[i][1], reverse=True)
    #     ios.sort(key=lambda i: kwarg[i][1], reverse=True)
    #     ids.sort(key=lambda i: kwarg[i][1], reverse=True)
    #     iS.sort(key=lambda i: kwarg[i][1], reverse=True)
    # else:  # cfca, on trie par score Origo
    ies.sort(key=lambda i: diffs_ergo[i], reverse=True)
    ios.sort(key=lambda i: diffs_origo[i], reverse=True)
    ids.sort(key=lambda i: diffs_dc[i], reverse=True)
    iS.sort(key=lambda i: diffs_origo[i], reverse=True)

    with open(file_out, 'w') as f:
        f.write('--- Règles Ergo : ---\n')
        f.write('\n'.join(['%s ;; (%sscore Ergo= %.3f)'
                           % (ruleName(header, indsX[i], indsY[i]),
                              kwName(kwarg[i]), diffs_ergo[i]) for i in ies]))
        f.write('\n\n--- Règles Origo : ---\n')
        f.write('\n'.join(['%s ;; (%sscore Origo= %.3f)'
                           % (ruleName(header, indsX[i], indsY[i]),
                              kwName(kwarg[i]), diffs_origo[i]) for i in ios]))
        f.write('\n\n--- Règles DC : ---\n')
        f.write('\n'.join(['%s ;; (%sscore DC= %.3f)'
                           % (ruleName(header, indsX[i], indsY[i]),
                              kwName(kwarg[i]), diffs_dc[i]) for i in ids]))
        f.write('\n\n\n--- Règles communes : ---\n')

        if isinstance(kwarg[0], str):
            f.write('\n'.join(['%s ;; (ergo=%.3f, origo=%.3f, dc=%.3f)'
                               % (ruleName(header, indsX[i], indsY[i]),
                                  diffs_ergo[i], diffs_origo[i], diffs_dc[i])
                               for i in iS]))
        else:
            f.write('\n'.join(['%s ;; (%s)'
                               % (ruleName(header, indsX[i], indsY[i]),
                                  kwName(kwarg[i])[:-2])
                               for i in iS]))
        f.write(bilan)
        f.close()

    print("Sauvegardé dans '%s'" % file_out)

    print("...terminé !")


def common_rule(version: int, regles: str):
    """ Trouve les règles inférées par plusieurs approches.

    Parameters
    ----------
    regles : str
        'Ergo', 'Origo', 'DC', ou 'communes'.
        Le type d'algorithme de la fonction origo() qu'on souhaite observer.
    """
    s = '------ %s v%i ------' % (regles, version)
    fichiers = ['Expe_agro_Zac_v%i TRIÉ' % version,
                'result_bicpams_additive_v%i_TRIÉ' % version,
                'result_bicpams_constant_v%i_TRIÉ' % version,
                'coron_rules_v%i' % version]
    noms = ['CFCA', 'bicpams_additive', 'bicpams_constant', 'coron']
    data = []
    for (i, fi) in enumerate(fichiers):
        with open('files/5_origo/' + fi + '_origo.txt', 'r') as f:
            data.append(f.readlines())
            a = 0
            while data[i][a] != '--- Règles %s : ---\n' % regles:
                a += 1
            b = a
            while data[i][b] != '\n':
                b += 1
            data[i] = set([data[i][x].split(' ;; ')[0]
                           for x in range(a + 1, b)])

    for i in range(len(data)):
        for j in range(i):
            inte = data[i].intersection(data[j])
            if inte:
                s += '\n%s / %s :\n  - %s\n' \
                        % (noms[i], noms[j], '\n  - '.join(inte))
    return s + '\n\n'


def common_rules():
    with open('files/common_rules.txt', 'w') as f:
        for r in ['Origo', 'Ergo', 'DC', 'communes']:
            for version in [1, 2]:
                f.write(common_rule(version, r))
        f.close()


##

dimY = 5  # ne pas toucher, paramètres correspondant à nos données
catcols = 2


def test():
    """ Sur les deux versions des indicateurs,
        on exécute chaque approche (séparées par un saut de ligne).
        Les lignes commentées sont longues d'exécution.

        Attention : entre les lignes bicpams_preprocess()
        et bicpams_postprocess(), il est nécessaire de générer le fichier .htm
        à partir du logiciel codes_sources/bicpams_4.0.4_gui.jar.
        Normalization : Column, Noise handler: Multitem,
        Coherency assumption : Constant ou Additive,
        Load data matrix : 'files/2a_bicpams/bicpams25_30.txt',
        Display solution : 'Biclusters (row/column names)'
                            Save -> 'result_bicpams_additive_v1.htm'

        """
    for version in [1, 2]:
            M, header = preprocess(version, verb=True)

            # CFCA_process(version)
            CFCA_postprocess(version, dimY)

            bicpams_preprocess(version)
            for coh in ['additive', 'constant']:
                bicpams_postprocess(version, coh, dimY)

            pypingen_binarisation(version)
            pypingen_plusoumoins(version)
            # pypingen_process()

            min_supp = 30
            min_conf = 60
            coron_binarisation(version)
            coron_mat2File(version)
            coron_process(version, min_supp, min_conf)
            coron_postprocess(version, dimY)

            globe_preprocess(version, catcols)
            globe_process(version)
            globe_postprocess(version)

            for source in [
                        ('CFCA', True),
                        ('CFCA', False),
                        ('bicpams', 'add'),
                        ('bicpams', 'cst'),
                        'coron'
                        ]:
                origo(version, source, dimY)

    common_rules()
t = time.time()
try:
    test()
finally:
    print('%.3f secondes.' % (time.time()-t))
